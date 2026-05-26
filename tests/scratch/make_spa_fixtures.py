"""Generate the 5 SPA Chrome MCP acceptance fixtures.

Per kaos-modules/docs/plans/2026-05-26-corpus-stress-residuals-S16-S22-and-spa-acceptance.md
Work item 1 scope:

- S19_lineitems.xlsx     openpyxl-produced workbook with a planted row
- S03_scanned.pdf         image-only PDF (bitmap font) carrying a token
- S20_cfr_memo.docx       memo citing 17 CFR § 240.10b-5(b) + the regulation text
- S16_pile/               5 files (PDF + DOCX + XLSX + HTML + JSON) each with a needle
- S22_cluster.docx        narrative with multi-metric (total rows + distinct clusters)

Outputs to /tmp/spa-acceptance/.
"""

from __future__ import annotations

import io
import json
from pathlib import Path

OUT = Path("/tmp/spa-acceptance")
OUT.mkdir(parents=True, exist_ok=True)


def _synth_pdf_text(text: str, title: str = "") -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    if title:
        c.setTitle(title)
    pw, ph = LETTER
    y = ph - 72
    for line in text.split("\n"):
        c.drawString(72, y, line)
        y -= 16
    c.showPage()
    c.save()
    return buf.getvalue()


def _synth_pdf_image(text: str) -> bytes:
    from PIL import Image, ImageDraw, ImageFont
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    img = Image.new("RGB", (1700, 2200), color="white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()
    y = 60
    for line in text.split("\n"):
        draw.text((60, y), line, fill="black", font=font)
        y += 28
    png_buf = io.BytesIO()
    img.save(png_buf, format="PNG")
    png_buf.seek(0)

    pdf_buf = io.BytesIO()
    c = canvas.Canvas(pdf_buf, pagesize=LETTER)
    pw, ph = LETTER
    c.drawImage(
        ImageReader(png_buf),
        x=0,
        y=0,
        width=pw,
        height=ph,
        preserveAspectRatio=True,
        mask="auto",
    )
    c.showPage()
    c.save()
    return pdf_buf.getvalue()


def _synth_docx(text: str, title: str = "") -> bytes:
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _synth_xlsx_openpyxl(rows: list[list], sheet: str = "Sheet1") -> bytes:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def write(name: str, data: bytes) -> Path:
    path = OUT / name
    path.write_bytes(data)
    print(f"  wrote {path}  ({len(data)} bytes)")
    return path


def main() -> None:
    print(f"[+] writing fixtures to {OUT}/")

    # --- S19: openpyxl-produced XLSX with a planted needle row
    print("[S19] openpyxl XLSX with planted needle")
    rows = [
        ["row_id", "description", "amount_usd"],
        ["LI-S19-9000", "Vendor onboarding", "2500.00"],
        ["LI-S19-9001", "Critical line item", "77777.77"],
        ["LI-S19-9002", "Catering", "612.50"],
    ]
    write("S19_lineitems.xlsx", _synth_xlsx_openpyxl(rows, sheet="LineItems"))

    # --- S03: image-only PDF requiring OCR
    print("[S03] scanned PDF requiring OCR (token: FALCON)")
    write(
        "S03_scanned.pdf",
        _synth_pdf_image(
            "INTERNAL MEMO\n"
            "Codename: FALCON\n"
            "Distribution: leadership only.\n"
            "Effective date: 2026-08-01\n"
        ),
    )

    # --- S20: memo citing 17 CFR § 240.10b-5(b) + the regulation text
    print("[S20] memo citing 17 CFR § 240.10b-5(b)")
    s20_memo = (
        "MEMORANDUM\n"
        "TO: Counsel\n"
        "FROM: Compliance\n"
        "RE: Anti-fraud — 17 CFR § 240.10b-5\n\n"
        "Please review the operative anti-fraud rule below. The full text of "
        "17 CFR § 240.10b-5(b) provides:\n\n"
        "It shall be unlawful for any person, directly or indirectly, by the use "
        "of any means or instrumentality of interstate commerce, or of the mails "
        "or of any facility of any national securities exchange, "
        "(b) to make any untrue statement of a material fact or to omit to state "
        "a material fact necessary in order to make the statements made, in the "
        "light of the circumstances under which they were made, not misleading.\n\n"
        "Counsel should flag any draft language that risks violating this rule."
    )
    write("S20_cfr_memo.docx", _synth_docx(s20_memo, title="Anti-Fraud Memo"))

    # --- S16: 5-format pile, each carrying a distinctive needle
    print("[S16] 5-format pile (PDF + DOCX + XLSX + HTML + JSON)")
    pile = OUT / "S16_pile"
    pile.mkdir(exist_ok=True)
    (pile / "pile-revenue.pdf").write_bytes(
        _synth_pdf_text(
            "REVENUE BRIEF\n\nFormat-pile PDF fact: revenue plug-figure $93.14M.\n\n"
            "All other line items confirmed to roll forward.",
            title="Revenue Brief",
        )
    )
    print(f"  wrote {pile / 'pile-revenue.pdf'}")
    (pile / "pile-counsel.docx").write_bytes(
        _synth_docx(
            "COUNSEL OF RECORD\n\n"
            "Format-pile DOCX fact: counsel-of-record is Hannah Brueggeman.\n\n"
            "Engagement letter executed on 2026-03-01.",
            title="Counsel Memo",
        )
    )
    print(f"  wrote {pile / 'pile-counsel.docx'}")
    (pile / "pile-lineitems.xlsx").write_bytes(
        _synth_xlsx_openpyxl(
            [
                ["row_id", "description", "amount_usd"],
                ["LI-PILE-9000", "Vendor onboarding", "2500.00"],
                ["LI-PILE-9001", "Pile-test critical line", "77777.77"],
                ["LI-PILE-9002", "Catering", "612.50"],
            ],
            sheet="LineItems",
        )
    )
    print(f"  wrote {pile / 'pile-lineitems.xlsx'}")
    (pile / "pile-launch.html").write_text(
        "<!DOCTYPE html>\n<html><body><h1>LAUNCH ANNOUNCEMENT</h1>"
        "<p>Format-pile HTML fact: marketing launch on 2026-08-14.</p>"
        "<p>Pre-orders open one week prior.</p></body></html>"
    )
    print(f"  wrote {pile / 'pile-launch.html'}")
    (pile / "pile-config.json").write_text(
        json.dumps(
            {
                "environment": "production",
                "release_token": "PILE-JSON-TOKEN-Z7Q",
                "feature_flags": {"format_pile": True},
            },
            indent=2,
        )
    )
    print(f"  wrote {pile / 'pile-config.json'}")

    # --- S22: cluster routing — multi-metric narrative (total rows + distinct clusters)
    print("[S22] cluster routing — labelled multi-metric")
    s22_doc = (
        "CLUSTER ROUTING REPORT\n\n"
        "Q4 2025 routing review. The full dataset contains 2486 total rows of "
        "transaction records sourced from the upstream operational warehouse.\n\n"
        "Each row was assigned to one of 24 distinct routing clusters based on "
        "destination zone + product class + handler queue. The 24 clusters are "
        "the policy units; the 2486 rows are the underlying records.\n\n"
        "Question for the analyst: how many routing clusters exist, AND what is "
        "the total transaction-record count? Both numbers are policy-relevant."
    )
    write("S22_cluster.docx", _synth_docx(s22_doc, title="Cluster Routing Report"))

    print(f"\n[+] done. fixtures in {OUT}/")


if __name__ == "__main__":
    main()
