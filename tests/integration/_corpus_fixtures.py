"""Dynamic in-memory fixture generators for the corpus stress suite.

Hermetic test fixtures (no network, no on-disk prefab corpora) for the
multi-tier corpus stress tests in ``test_corpus_stress_suite.py``. Each
helper returns *raw bytes* with proper magic-byte signatures so the
content-type detector (``kaos_nlp_core.content_type.detect``) sees the
real format regardless of the filename we hang on it.

Capabilities:

- ``synth_pdf_text(text, *, title="")``: text-layer PDF via reportlab.
- ``synth_pdf_image(text)``: rasterized-text PDF — embeds a PIL-rendered
  PNG so the only way to recover text is OCR. Used to stress the OCR
  branch of the upload pipeline (S03).
- ``synth_docx(text, *, title="")``: native DOCX via python-docx.
- ``synth_xlsx(rows, *, sheet_name="Sheet1")``: native XLSX via openpyxl.
- ``synth_html(text, *, title="")``: minimal but valid HTML5 doc.
- ``synth_text(text)``: UTF-8 bytes for ``.txt`` payloads.
- ``synth_json(payload)``: pretty-printed UTF-8 JSON bytes.
- ``wrong_extension(content_bytes, advertised_ext)``: (bytes, filename)
  where the filename uses the advertised extension while the bytes keep
  the original magic signature. The (filename, bytes) pair is the
  spoof: the content-type detector should report the *real* format from
  the bytes; the agent's tool stack must not blindly trust the
  extension. Tests S02 and S10 ride on this.

- ``SynthDoc`` dataclass: filename + bytes + real MIME + needle/non-needle
  flag + (optional) the planted ground-truth fact.

- ``synth_corpus(n_docs, n_needles, needle_facts, *, ...)``: produces an
  N-doc corpus with planted needles at deterministic positions, optional
  topical clustering, and optional wrong-extension spoofing. Distractor
  prose is generated from a small canned vocabulary (no network at test
  time) — predictable enough that BM25 has signal but verbose enough
  that a needle isn't a trivial top-1.

- ``write_corpus_to_vfs(docs, runtime, *, session_id)``:

  1. Writes each doc's bytes to ``sessions/{session_id}/files/{filename}``
     on the runtime VFS, mirroring the SPA's layout
     (``kaos-ui/examples/single-user-chat/backend/app/services/uploads.py``).
  2. Registers each as a ``BODY`` artifact via
     ``runtime.artifacts.create_from_path``.

  Returns the list of ``ArtifactManifest`` so tests can assert
  artifact_ids landed.

- ``hydrate_corpus_into_memory(docs, memory)``: convenience for tests
  that need the documents in ``MemoryType.DOCUMENTS`` (so the agent's
  ``triage_corpus`` / ``search_memory`` paths fire). Plants the full
  body text into the section and tags each item with the same
  ``filename`` / ``vfs_path`` metadata the SPA uses.

Hermeticity rule: every generator is deterministic given its inputs +
seed. No HTTP / DNS / clock dependence. The only non-determinism allowed
in the stress suite proper is the LLM call.
"""

from __future__ import annotations

import io
import json
import random
from dataclasses import dataclass
from typing import TYPE_CHECKING, Any, Literal

from kaos_agents.memory.session import SessionMemory
from kaos_agents.types.memory import MemoryType

if TYPE_CHECKING:
    from collections.abc import Sequence

    from kaos_core.artifacts.models import ArtifactManifest
    from kaos_core.registry.container import KaosRuntime


# ---------------------------------------------------------------------------
# Single-format generators
# ---------------------------------------------------------------------------


def synth_pdf_text(text: str, *, title: str = "") -> bytes:
    """Build a text-layer PDF from ``text`` via reportlab.

    The output has a real text layer (pypdfium2 / pdfminer can extract
    it without OCR). Long bodies get wrapped to ~88-char lines and
    paginated automatically.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
        title=title or "Synthetic Document",
    )
    styles = getSampleStyleSheet()
    story: list[Any] = []
    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))
    # Split on blank lines into paragraphs; preserve single-line breaks
    # inside a paragraph via <br/>.
    for raw_para in text.split("\n\n"):
        para = raw_para.strip()
        if not para:
            continue
        # reportlab uses & < > as markup; escape them defensively.
        safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe = safe.replace("\n", "<br/>")
        story.append(Paragraph(safe, styles["BodyText"]))
        story.append(Spacer(1, 6))
    doc.build(story)
    return buf.getvalue()


def synth_pdf_image(text: str) -> bytes:
    """Build a PDF whose only content is a rasterized image of ``text``.

    No text layer. Recovering ``text`` requires OCR — exactly the path
    we want to stress in S03. The image is generated in-memory via PIL
    and embedded as a single full-page PNG.
    """
    from PIL import Image, ImageDraw, ImageFont
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    width_px, height_px = 1700, 2200  # ~200 dpi @ Letter
    img = Image.new("RGB", (width_px, height_px), color="white")
    draw = ImageDraw.Draw(img)
    # default bitmap font — installed everywhere PIL is. Looks like a
    # scanned typewriter page, which is what we want for OCR exercise.
    font = ImageFont.load_default()
    # Wrap to ~80 chars.
    margin = 60
    line_height = 28
    y = margin
    for raw_line in text.split("\n"):
        chunks = [raw_line[i : i + 80] for i in range(0, max(1, len(raw_line)), 80)]
        for chunk in chunks:
            draw.text((margin, y), chunk, fill="black", font=font)
            y += line_height
            if y > height_px - margin:
                break
        if y > height_px - margin:
            break
    png_buf = io.BytesIO()
    img.save(png_buf, format="PNG")
    png_buf.seek(0)

    pdf_buf = io.BytesIO()
    c = canvas.Canvas(pdf_buf, pagesize=LETTER)
    page_w, page_h = LETTER
    # Embed the raster image so it fills the page.
    from reportlab.lib.utils import ImageReader

    c.drawImage(
        ImageReader(png_buf),
        x=0,
        y=0,
        width=page_w,
        height=page_h,
        preserveAspectRatio=True,
        anchor="c",
        mask="auto",
    )
    c.showPage()
    c.save()
    return pdf_buf.getvalue()


def synth_docx(text: str, *, title: str = "") -> bytes:
    """Build a DOCX from ``text`` via python-docx.

    Each blank-line-separated chunk in ``text`` becomes a paragraph.
    """
    from docx import Document

    document = Document()
    if title:
        document.add_heading(title, level=1)
    for raw_para in text.split("\n\n"):
        para = raw_para.strip()
        if not para:
            continue
        document.add_paragraph(para)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def synth_xlsx(rows: list[list[str]], *, sheet_name: str = "Sheet1") -> bytes:
    """Build an XLSX from ``rows`` (list of row-lists) via openpyxl.

    The first row is treated as the header. Empty cells are written as
    empty strings, not None, so downstream parsers don't have to special-
    case Nones.
    """
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    if ws is None:
        # openpyxl always provides an active sheet for a new Workbook,
        # but ty's narrow signature doesn't know that.
        ws = wb.create_sheet()
    ws.title = sheet_name[:31]  # Excel hard cap on sheet name length
    for row in rows:
        ws.append([cell if cell is not None else "" for cell in row])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def synth_html(text: str, *, title: str = "") -> bytes:
    """Build a minimal valid HTML5 doc whose body contains ``text``.

    Paragraphs are split on blank lines (mirrors synth_pdf_text /
    synth_docx). HTML special chars in ``text`` are escaped.
    """
    import html

    paragraphs = []
    for raw_para in text.split("\n\n"):
        para = raw_para.strip()
        if not para:
            continue
        paragraphs.append(f"<p>{html.escape(para)}</p>")
    body = "\n".join(paragraphs)
    safe_title = html.escape(title) if title else "Synthetic Document"
    return (
        f"<!doctype html>\n"
        f'<html lang="en">\n'
        f'<head><meta charset="utf-8"><title>{safe_title}</title></head>\n'
        f"<body>\n{body}\n</body>\n</html>\n"
    ).encode()


def synth_text(text: str) -> bytes:
    """Plain UTF-8 text."""
    return text.encode("utf-8")


def synth_json(payload: dict[str, Any]) -> bytes:
    """Pretty-printed UTF-8 JSON."""
    return json.dumps(payload, indent=2, sort_keys=True).encode("utf-8")


# ---------------------------------------------------------------------------
# Extension spoofing
# ---------------------------------------------------------------------------


def wrong_extension(content_bytes: bytes, advertised_ext: str) -> tuple[bytes, str]:
    """Return ``(bytes, filename)`` where ``filename`` carries
    ``advertised_ext`` but ``bytes`` keep their original magic-byte
    signature.

    Example::

        body, name = wrong_extension(synth_pdf_text("hi"), ".txt")
        # name == "fixture-XXXXXXXX.txt"
        # body[:5] == b"%PDF-"  (still a PDF)

    Test S02 / S10 rely on this exact mismatch to verify that the
    upload pipeline content-sniffs bytes rather than trusting the
    extension.
    """
    if not advertised_ext.startswith("."):
        advertised_ext = "." + advertised_ext
    # Use a content-hash prefix so the same input pair always names the
    # same file — keeps the corpus deterministic.
    import hashlib

    digest = hashlib.sha256(content_bytes).hexdigest()[:8]
    return content_bytes, f"fixture-{digest}{advertised_ext}"


# ---------------------------------------------------------------------------
# Synthetic corpus
# ---------------------------------------------------------------------------


@dataclass(frozen=True, slots=True)
class SynthDoc:
    """A single fixture document for the stress corpus.

    Attributes
    ----------
    filename
        On-disk name (e.g. ``"doc-042.pdf"``). For spoofed entries this
        ends in the *wrong* extension on purpose — the body still has
        the real magic bytes.
    bytes
        Raw file bytes with real magic signature.
    mime
        The *real* MIME type as derivable from the bytes (independent of
        the filename extension).
    is_needle
        True when this doc carries one of the planted facts.
    needle_fact
        The exact planted phrase (e.g. ``"Project Oryx codename:
        7K-FALCON-2026"``) when ``is_needle`` is True; otherwise None.
    """

    filename: str
    bytes: bytes
    mime: str
    is_needle: bool
    needle_fact: str | None = None


# Tiny canned vocabulary for distractor prose. Keep it boring but
# topical so BM25 can rank — pure lorem-ipsum has no signal and the
# triage step would behave randomly. Each "cluster" is a topic bucket
# so cluster-routing scenarios (S08) have something to route over.
_CLUSTER_VOCAB: dict[str, list[str]] = {
    "finance": [
        "The Q3 earnings report shows a 12% revenue increase year-over-year.",
        "Operating margin compressed from 18.4% to 15.7% on cost inflation.",
        "Free cash flow guidance for fiscal 2026 was reaffirmed at $1.2B.",
        "Working capital efficiency improved due to a faster A/R cycle.",
        "Capex intensity dropped to 4.2% of revenue versus 5.1% prior.",
        "The board approved a $500M share repurchase program in May.",
        "EBITDA margin held flat as logistics savings offset wage growth.",
        "Net leverage remained below 2x with $2.4B of liquidity.",
    ],
    "legal": [
        "Counsel reviewed the indemnification carveouts in Schedule 4.2.",
        "The non-compete clause restricts engagement for twenty-four months.",
        "Section 7.3 of the master services agreement governs IP assignment.",
        "Termination for convenience requires sixty days written notice.",
        "Mutual non-disclosure obligations survive for three years post-term.",
        "Governing law is Delaware; venue lies exclusively in Wilmington.",
        "Liquidated damages cap at the lesser of $1M or trailing twelve fees.",
        "Force majeure events suspend performance but not payment.",
    ],
    "engineering": [
        "The microservice exposes a gRPC interface on port 8443.",
        "Latency p99 dropped from 240ms to 90ms after the cache redesign.",
        "Postgres autovacuum was tuned to avoid table bloat on writes.",
        "The shard key uses (tenant_id, created_at) for time-locality.",
        "Connection pooling was migrated from pgBouncer to RDS Proxy.",
        "Retry budgets are bounded to 3 attempts with 250ms exponential backoff.",
        "OpenTelemetry traces flow through Tempo with 14-day retention.",
        "The CI pipeline runs unit + integration tiers on every PR.",
    ],
    "medical": [
        "Patient cohort A showed a 38% response rate at 12-week follow-up.",
        "Adverse event grades stayed below CTCAE grade 3 throughout the arm.",
        "PK/PD modeling indicates a 14-hour half-life in healthy volunteers.",
        "The dose-escalation phase enrolled 24 patients across three cohorts.",
        "Pre-specified secondary endpoints reached statistical significance.",
        "Concomitant medication restrictions excluded strong CYP3A4 inhibitors.",
        "Stratification factors included age, ECOG status, and prior therapy.",
        "Cold-chain handling required transport at 2-8C per protocol.",
    ],
    "regulatory": [
        "The 2026 SEC climate disclosure rule applies to large filers first.",
        "FAR Part 12 governs commercial-item acquisitions for the agency.",
        "The HIPAA security rule mandates encryption of PHI in transit.",
        "Notification timelines under GDPR Article 33 run for 72 hours.",
        "Section 404 of Sarbanes-Oxley requires management ICFR attestation.",
        "FDA 510(k) submissions reference predicate devices for clearance.",
        "OFAC sanctions screening must run against the consolidated list.",
        "FERPA permits directory information disclosure absent opt-out.",
    ],
}


def _generate_distractor_prose(
    rng: random.Random,
    *,
    cluster: str,
    paragraphs: int,
) -> str:
    """Stitch ``paragraphs`` worth of canned vocabulary into a doc body."""
    pool = _CLUSTER_VOCAB.get(cluster, _CLUSTER_VOCAB["legal"])
    out: list[str] = []
    for _ in range(paragraphs):
        # 3-6 sentences per paragraph drawn with replacement.
        n_sent = rng.randint(3, 6)
        sentences = [rng.choice(pool) for _ in range(n_sent)]
        out.append(" ".join(sentences))
    return "\n\n".join(out)


def _real_mime_for_ext(ext: str) -> str:
    return {
        ".pdf": "application/pdf",
        ".docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ".xlsx": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        ".html": "text/html",
        ".txt": "text/plain",
        ".json": "application/json",
    }.get(ext.lower(), "application/octet-stream")


def _render_for_type(
    file_type: str,
    *,
    body: str,
    title: str,
    needle_fact: str | None,
) -> tuple[bytes, str, str]:
    """Render ``body`` (with optional planted needle) into the chosen
    file format. Returns ``(bytes, ext, mime)``.
    """
    # If we have a needle, prepend it so retrievers see it near the top
    # AND a generous head-truncation summary captures it.
    if needle_fact:
        body = needle_fact + "\n\n" + body

    if file_type == "pdf":
        return synth_pdf_text(body, title=title), ".pdf", _real_mime_for_ext(".pdf")
    if file_type == "docx":
        return synth_docx(body, title=title), ".docx", _real_mime_for_ext(".docx")
    if file_type == "xlsx":
        # XLSX needs row data: turn paragraphs into rows with one cell each.
        rows: list[list[str]] = [["section", "text"]]
        for i, para in enumerate(body.split("\n\n")):
            rows.append([f"p{i + 1}", para.strip()])
        return synth_xlsx(rows, sheet_name="Body"), ".xlsx", _real_mime_for_ext(".xlsx")
    if file_type == "html":
        return synth_html(body, title=title), ".html", _real_mime_for_ext(".html")
    if file_type == "text":
        return synth_text(body), ".txt", _real_mime_for_ext(".txt")
    if file_type == "json":
        payload = {
            "title": title,
            "body": body,
            "needle": needle_fact,
        }
        return synth_json(payload), ".json", _real_mime_for_ext(".json")
    raise ValueError(f"unsupported file_type: {file_type!r}")


def synth_corpus(
    n_docs: int,
    n_needles: int,
    needle_facts: list[str],
    *,
    seed: int = 0,
    cluster_topics: list[str] | None = None,
    file_types: Sequence[Literal["pdf", "docx", "xlsx", "html", "text", "json"]] = (
        "pdf",
        "docx",
        "html",
        "text",
    ),
    wrong_extension_rate: float = 0.0,
    paragraphs_per_doc: int = 6,
) -> list[SynthDoc]:
    """Generate ``n_docs`` deterministic synthetic documents.

    Parameters
    ----------
    n_docs
        Total documents in the corpus (needles + distractors).
    n_needles
        How many of the docs carry a planted needle. Must be
        ``<= len(needle_facts)`` and ``<= n_docs``.
    needle_facts
        The planted ground-truth phrases — one per needle. Each phrase
        should be exact-substring-checkable in the agent's response
        (no fuzzy matching on the test side).
    seed
        RNG seed for deterministic placement / file-type assignment.
    cluster_topics
        Optional list of topic labels. Non-needle docs are distributed
        across these clusters round-robin so cluster-routing scenarios
        (S08) have a topical signal. Defaults to ``["finance", "legal",
        "engineering", "medical", "regulatory"]``.
    file_types
        Which file formats may be assigned to docs. Default mix covers
        the file-loading hot path (PDF + DOCX + HTML + plain text).
    wrong_extension_rate
        Fraction of docs (in [0, 1]) whose filename will end in the
        wrong extension — bytes keep their real magic signature.
        Defaults to 0 (off).
    paragraphs_per_doc
        Approximate body length. Real production corpora are much
        larger; we keep test docs short so the suite still runs in CI
        budget but long enough to give BM25 signal.

    Returns
    -------
    list[SynthDoc]
        Stable order: needles first (at deterministic spread positions),
        distractors filling the rest.
    """
    if n_needles > n_docs:
        raise ValueError(f"n_needles={n_needles} > n_docs={n_docs}")
    if n_needles > len(needle_facts):
        raise ValueError(f"n_needles={n_needles} > len(needle_facts)={len(needle_facts)}")

    rng = random.Random(seed)
    clusters = cluster_topics or ["finance", "legal", "engineering", "medical", "regulatory"]

    # Deterministic needle positions: evenly spread through the corpus
    # so a "first 10" hack can't accidentally find them all.
    if n_needles == 0:
        needle_positions: set[int] = set()
    elif n_needles == 1:
        # Single-needle scenarios: park it near the middle so BM25 has
        # to actually retrieve, not get-by-position.
        needle_positions = {n_docs // 2}
    else:
        step = max(1, n_docs // n_needles)
        needle_positions = {min(n_docs - 1, i * step) for i in range(n_needles)}

    docs: list[SynthDoc] = []
    needle_iter = iter(needle_facts[:n_needles])
    for i in range(n_docs):
        cluster = clusters[i % len(clusters)]
        is_needle = i in needle_positions
        needle_fact = next(needle_iter) if is_needle else None
        title = f"Document {i + 1:04d} ({cluster})"
        body = _generate_distractor_prose(
            rng,
            cluster=cluster,
            paragraphs=paragraphs_per_doc,
        )
        # Round-robin file type so we always exercise the full mix.
        file_type = file_types[i % len(file_types)]
        data, ext, mime = _render_for_type(
            file_type,
            body=body,
            title=title,
            needle_fact=needle_fact,
        )
        # Optional spoof.
        spoof = wrong_extension_rate > 0.0 and rng.random() < wrong_extension_rate
        if spoof:
            wrong_ext_pool = [e for e in (".pdf", ".docx", ".html", ".txt") if e != ext]
            advertised = rng.choice(wrong_ext_pool)
            _spoof_bytes, spoof_name = wrong_extension(data, advertised)
            filename = spoof_name
        else:
            filename = f"doc-{i + 1:04d}{ext}"
        docs.append(
            SynthDoc(
                filename=filename,
                bytes=data,
                mime=mime,  # the REAL mime, regardless of the on-disk extension
                is_needle=is_needle,
                needle_fact=needle_fact,
            )
        )
    return docs


# ---------------------------------------------------------------------------
# VFS / SessionMemory plumbing
# ---------------------------------------------------------------------------


async def write_corpus_to_vfs(
    docs: list[SynthDoc],
    runtime: KaosRuntime,
    *,
    session_id: str,
) -> list[ArtifactManifest]:
    """Write each doc's bytes to ``sessions/{session_id}/files/{filename}``
    and register as a BODY artifact.

    Mirrors the SPA upload pipeline
    (``kaos-ui/examples/single-user-chat/backend/app/services/uploads.py``)
    so the agent's tool surface resolves bare filenames at the same
    on-disk location it would in production.

    Returns the list of ``ArtifactManifest`` in the same order as
    ``docs`` so tests can spot-check ``artifact_id`` / ``path`` /
    ``mime_type`` against ``SynthDoc.mime``.
    """
    manifests: list[ArtifactManifest] = []
    for doc in docs:
        vfs_path = f"sessions/{session_id}/files/{doc.filename}"
        await runtime.vfs.write(vfs_path, doc.bytes)
        manifest = await runtime.artifacts.create_from_path(
            vfs_path,
            context_id=session_id,
            session_id=session_id,
            name=doc.filename,
            # Pass the *real* mime so an artifact-consumer that trusts
            # the manifest gets the bytes-truthful type, while the
            # on-disk filename keeps the (possibly spoofed) extension.
            mime_type=doc.mime,
        )
        manifests.append(manifest)
    return manifests


def hydrate_corpus_into_memory(
    docs: list[SynthDoc],
    memory: SessionMemory,
    *,
    session_id: str | None = None,
) -> None:
    """Plant each doc into ``MemoryType.DOCUMENTS`` with SPA-shape metadata.

    The agent's ``triage_corpus`` / ``search_memory`` paths read from
    ``DOCUMENTS``, so a test that just wrote bytes to the VFS without
    also seeding memory would force the agent to discover the corpus
    via ``kaos-core-vfs-list`` first — which is a different test (we
    cover that separately).

    For text-extractable formats (text, html, json) we plant the actual
    body so BM25 has signal. For binary formats (pdf, docx, xlsx) we
    plant a metadata headline + the needle fact when present; the agent
    must call a parser tool to get the full body, matching the
    production retrieval contract.
    """
    for doc in docs:
        if doc.mime.startswith("text/") or doc.mime == "application/json":
            content = doc.bytes.decode("utf-8", errors="replace")
        else:
            # Mirror the SPA's DOCUMENTS-section "headline" shape.
            headline_parts = [
                f"filename: {doc.filename}",
                f"vfs_path: sessions/{session_id or 'NA'}/files/{doc.filename}",
                f"size_bytes: {len(doc.bytes)}",
                f"content_type: {doc.mime}",
            ]
            if doc.needle_fact:
                headline_parts.append(f"summary: {doc.needle_fact}")
            content = " | ".join(headline_parts)
        memory.add(
            MemoryType.DOCUMENTS,
            content,
            metadata={
                "uri": f"file:{doc.filename}",
                "filename": doc.filename,
                "vfs_path": (
                    f"sessions/{session_id}/files/{doc.filename}" if session_id else doc.filename
                ),
                "mime_type": doc.mime,
            },
        )
